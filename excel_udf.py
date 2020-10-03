			#Excel File Creator 


def excel(filename):

	#Imports	
	import xlsxwriter
	import docx
	from docx.api import Document
	import readDocx
	import re
	import unicodedata
	
	#Progress Bar
	#from progressbar import ProgressBar
	#pbar = ProgressBar()

	#Determining the names of files as list 
	filename_ = filename[0:len(filename)-5]
	#Number of files
	num_files = list(range(len(filename)))
	
	#Creating a .xml file for every .docx
	#Initializing table data sync 
	date = []
	for i in num_files:
		for find in re.finditer('Log', filename):
			if find:
				span = find.span()
				leng = len(filename)
				dates = filename[span[1]+1:leng-5]
				date.append(dates)

	workbook = xlsxwriter.Workbook('C:/Users/dalrympm/Documents/Coding/Files/TMO_Spread_Sheets/'+str(filename_)+"_spreadsheet.xlsx")

	worksheet = workbook.add_worksheet()

	# Create a format to use in title cells.
	cell_format_title = workbook.add_format({
     	'font_size': '22',
     	'bg_color': '#f3a030',
     	'underline':'true',
     	'border': 2,
     	'align': 'center',
     	'valign': 'vcenter',
     	'bold': 'True' })
    

	# Create a format to use in populated cells.
	cell_format_pop = workbook.add_format({
    	'bold': 1,
    	'border': 1,
    	'align': 'center',
    	'font_color': '#000000',
    	'font_size': '24',
    	'valign': 'vcenter'})

	# Create a format to use in block cells.
	cell_format_info = workbook.add_format({
	    'bold': 1,
	    'border': 1,
	    'align': 'center',
	    'font_color': '#000000',
	    'font_size': '34',
	    'valign': 'vcenter',
	    'text_wrap': 'vjustify'})


	#Setting widths of columns 
	worksheet.set_column('A:A', 62)
	worksheet.set_column('B:B', 64)
	worksheet.set_column('C:C', 60)
	worksheet.set_column('D:D', 60)
	worksheet.set_column('E:E', 62)
	worksheet.set_column('F:F', 70)
	worksheet.set_column('G:G', 68)
	worksheet.set_column('H:H', 70)
	worksheet.set_column('I:I', 64)
	worksheet.set_column('J:J', 64)
	worksheet.set_column('K:K', 74)
	worksheet.set_column('L:L', 60)
	worksheet.set_column('M:M', 64)
	worksheet.set_column('N:N', 64)
	worksheet.set_column('O:O', 70)
	worksheet.set_column('Q:Q', 76)

	#Setting column titles
	worksheet.write('A6', 'Delay #',cell_format_title)
	worksheet.write('B6', 'Date',cell_format_title)
	worksheet.write('C6', 'DIN',cell_format_title)
	worksheet.write('D6', 'Year',cell_format_title)
	worksheet.write('E6', 'Series',cell_format_title)
	worksheet.write('F6', 'Reactor Face',cell_format_title)
	worksheet.write('G6', 'Latice Site',cell_format_title)
	worksheet.write('H6', 'OPN Affected',cell_format_title)
	worksheet.write('I6', 'Start Time',cell_format_title)
	worksheet.write('J6', 'End Time',cell_format_title)
	worksheet.write('K6', 'Critical Path (Y/N)',cell_format_title)
	worksheet.write('L6', 'Duration',cell_format_title)
	worksheet.write('M6', 'Delay Type',cell_format_title)
	worksheet.write('N6', 'Issue Group',cell_format_title)
	worksheet.write('O6', 'Tool Impacted',cell_format_title)
	worksheet.write('P6', 'Issue Description',cell_format_title)
	worksheet.write('Q6', 'Cumulative Delay(hrs)',cell_format_title)


	#Setting row heights, **row numbers pushed forward 1(ie, 1=2)**
	worksheet.set_row(0, 30)
	worksheet.set_row(1, 15)
	worksheet.set_row(2, 15)
	worksheet.set_row(3, 0)
	worksheet.set_row(4, 0)
	worksheet.set_row(5, 80)
	for i in range(10000):
		worksheet.set_row(i+6,120)
	exit

	# Create a format to use in the merged range.
	merge_format = workbook.add_format({
    	'bold': 1,
    	'border': 1,
    	'align': 'center',
   	 'valign': 'vcenter',
    	'fg_color': ' #FFA500'})
    	
    

										#WORD DOCX PARSER

	document = Document(filename)

		#Finds and counts all tables in .docx
	tables_counter = document.tables

		#Make list of how many tables in .docx
	table_quantity = len(tables_counter)
	table_list = list(range(table_quantity))

		#Variables to populate
	data = []
	t_titles = []
	sub_sections = []
	sub_sections_ref = []
		
		#Loop through each TABLE and save text
	for x in table_list:
			#Calling on Word .docx Table
		table = document.tables[x]
		keys = None

		for i, row in enumerate(table.rows):
    			text = (cell.text for cell in row.cells)

    			if i == 0:
					#Keys are the title cells of each table
        			keys = tuple(text)			
        			continue
					#Saving Keys not working right now_ still(no further explanation
			
				#Saving text as data
    			row_data = tuple(text)
    			data.append(row_data)

			#Amount of subsections
		sub_section = len(data)
			#Subsections refural list
		sub_sections_ref.append(sub_section)	
	
			#Sub_sections is the amount of sub section each table has 
		if x < 1:
			sub_sections.append(sub_section)	
	
		if x >= 1:
			sub_section = len(data) - sub_sections_ref[x-1]
			sub_sections.append(sub_section)	
	exit

				#Excel cell population
	count_num = []
	main_sites = []	
	count = 6
	coun = 0		 
	for i in table_list:
			#Information in First Table
	
		find = re.match("\AEast|West|RAB|RFRISA|RCC", data[i][0])
		subsub_section = list(range(1,len(data[i])))
	
			#Reactor face cell popuation
		if find:		
			span = find.span()
		
				#Reactor Face cell
			r_face = data[i][0][span[0]:span[1]]
			count += 1
			count_num.append(count)
		
				#Issue Description cell
			for x in subsub_section:
				count += 1
				coun += 1
				count_num.append(count)
				issue_d = data[i][x]
		
				loc_x = str('P'+str(count))
				loc_i = str('F'+str(count))
				loc_d = str('A'+str(count))

				worksheet.write(loc_x,issue_d,cell_format_info)
				worksheet.write(loc_i,r_face,cell_format_pop)
				worksheet.write(loc_d,coun,cell_format_info)
			exit	
		
 	
		else:

				#Turn tuple data into string array set and replace the line indicators
			me = str(data[i])
			#Had to replace the '\\n' indicator to 'mmm' because regEx was not identifying the indicators 	
			me = me.replace("\\n","mmm")

				#Variables to populate
			cell_data = []
			reactor_face = []
			span_list = []
			s_list = []
			if count_num == []:
				count = 6 
			else:
				count = max(count_num)
			#Find Selected words from Headings	
			for new_data in re.finditer("mmmEastmmm|mmmEast mmm|mmmEASTmmm|mmmEAST:mmm|mmmEAST mmm|mmmWestmmm|mmmWest mmm|mmmWESTmmm|mmmWEST:mmm|mmmWEST mmm", me):
				if new_data:
					span = new_data.span()
					span_list.append(span)
		
						#Reactor Face cell
					r_face = me[span[0]+3:span[1]-3] 
					reactor_face.append(r_face)
					s = list(range(len(span_list)))
					s_list.append(s)
				exit

			s_list2 = list(range(len(s_list)))
			sub_para = []
		
			#Locate how many sub-paragraphs are in each table
			for x in s_list2:
			
 
				if x == max(s_list2):
					count +=1
					sub_paragraphs = me[span_list[x][1]-3:len(me)]
					sub_para.append(str(sub_paragraphs)) 			

				else:
					sub_paragraphs = me[span_list[x][1]-3:span_list[x+1][0]]
					sub_para.append(str(sub_paragraphs))
				exit	
		
		
			span2_list = []
			s2_list = []
			sub_para_str = str(sub_para)

			#Split all the sub-paragraph data into sentences
			for sentence in re.finditer("m{3,}", sub_para_str):
				if sentence: 
					span2 = sentence.span()		
					span2_list.append(span2)
					s2 = list(range(len(span2_list)))
					s2_list.append(s2)
				exit
		
			s2_list2 = list(range(len(s2_list)))
			count_2 = 0		
			loc = []
			col_size = []
			sub_sentences = []
			
			#Spilt all senteces into separate strings 
			for z in s2_list2:	
			
				if z == max(s2_list2):

					sub_sentence = sub_para_str[span2_list[z][1]:len(me)]
					sub_sentences.append(sub_sentence)
			
				else:
					sub_sentence = sub_para_str[span2_list[z][1]:span2_list[z+1][0]]
					sub_sentence = sub_sentence.replace(u'\\xa0', u'')
					sub_sentences.append(sub_sentence)

			#Search for line punctures  
					if re.search("', '|', \"", sub_sentence):

							count += 1
							count_2 += 1						

					else:
						count += 1
						coun += 1

			#List each cell
						loc_i = str('F'+str(count))
						loc_x = str('P'+str(count))
						loc_d = str('A'+str(count))
					
						worksheet.write(loc_i,reactor_face[count_2],cell_format_pop)
						worksheet.write(loc_x,sub_sentence,cell_format_info)
						worksheet.write(loc_d,coun,cell_format_info)


						length = len(sub_sentence)
						col_size.append(length)
						worksheet.set_column('P:P', max(col_size)-100)
			#Variables to populate
						location =[]
						times = [] 

			#Find the time and date 
						for time in re.finditer("\d{1,2}:\d\dam|\d{1,2}:\d\dpm|\d{1,2}:\d\d|\d{1,2} oclock|\d{4}h|\d{1,2}noon", sub_sentence):
						
							if time:
								location.append(z)
								span3 = time.span()
								time = sub_sentence[span3[0]:span3[1]]	
								times.append(time)
								if len(location) == 2:
								
			#Populate cell with time&date data 
									loc_t = str('J'+str(count))
									worksheet.write(loc_t,time,cell_format_info)

							else:	
								loc_t = str('I'+str(count))
								worksheet.write(loc_t,time,cell_format_info)

						exit


			#Variables to populate
					location = []
					latice_sites = []
					
			#Find latice sites  
					for latice_site in re.finditer("[A-Z]{1}[0-9]{1,2} |[A-Z]{1}[0-9]{1,2}| [A-Z]{1}[0-9]{2} |[A-Z]{1}[0-9]{2},|[A-Z]{1}[0-9]{1},|[A-Z]{1}[0-9]{1,2}-|[A-Z]{1}[0-9]{1,2}-|[A-Z]{1}[0-9]{2}W|[A-Z]{1}[0-9]{1,2}E|[A-Z]{1}[0-9]{2}W:|[A-Z]{1}[0-9]{2}E:", sub_sentence):
	
						if latice_site:
							location.append(z)
							span4 = latice_site.span()
							latice_site = sub_sentence[span4[0]:span4[1]]
							latice_sites.append(latice_site)
							
							
							if len(location) == 1:
								main_sites.append(latice_site)
								loc_s = str('G'+str(count))
								worksheet.write(loc_s,latice_site,cell_format_info)
								#print(latice_site,"'First Site'")

							else:
								main = main_sites
								main.append(latice_site)
								cell_text = ','.join(main)
								#print(latice_site,"'More Site'")
								#print(cell_text)
								
			#Populate cell with latice sites data
								loc_c = str('G'+str(count))
								worksheet.write(loc_c,cell_text,cell_format_info)

	
			#Find OPN(###) 
						for opn in re.finditer("opn\d{3}", sub_sentence):
							if opn:
								span = opn.span()
								opn = sub_sentence[span[0]:span[1]]	
			#Populate cell with OPN(###) data	
								loc_o = str('H'+str(count))
								worksheet.write(loc_o,opn,cell_format_info)

			#Find Critical_path_delay(Y/N) sites 
						for crit_path in re.finditer("Critical path delay ~", sub_sentence):
							if crit_path:
								ans = 'Y'
								loc_c = str('K'+str(count))
								worksheet.write(loc_c,ans,cell_format_info)

			#Variables to populate
				count_lists = []
				count_numb = [] 
				main_sites = []		
				
			#Find Dates(MM/DD/YY) 		
				for date in re.finditer(r"\w+\s\d+\s\d{4}|\w+\s\d+th|\w+\s\d+\sDay|\w+\s\d+\sday|\w+\s\d+\s(Day)|\w+\s\d+\s(day)|\w+\s\d+\sDAY|\w+\s\d+\sNight|\w+\s\d+\snight|\w+\s\d+\s(Night)|\w+\s\d+\s(night)|\w+\s\d+\sNIGHT|\w+\s\d[1,2],",sub_sentence):

					if date:
						span = date.span()
						
			#Find Dates(/DD)		
					for day in re.finditer(r"\w+\s\d+\s|\w+\s\d+,|\w+\s\d+t", sub_sentence):
						if day:
							span2 = day.span()
							date = sub_sentence[span2[0]:span2[1]-1]
							count_numb.append(count)
							loc_d = str('B'+str(count))
							worksheet.write(loc_d,date,cell_format_info)

			#Find Dates(/YY)
					for year in re.finditer('\d{4}', sub_sentence):
						if year:
							span3 = year.span()
							year = sub_sentence[span3[0]:span3[1]]
							loc_y = str('D'+str(count)) 
							worksheet.write(loc_y,year,cell_format_info)

			#Find Dates(Nights)
					for shift_n in re.finditer('Night|night|NIGHT', sub_sentence):
						if shift_n:
							shift = "N"
							loc_s = str('C'+str(count)) 
							worksheet.write(loc_s,shift,cell_format_info)					

			#Find Dates(Day)
					for shift_d in re.finditer('Day|day|DAY', sub_sentence):
						if shift_d:
							shift = "D"
							loc_s = str('C'+str(count)) 
							worksheet.write(loc_s,shift,cell_format_info)


		exit				

				
#Same procedures for more than two table  
	if len(sub_sections) <= 20:

			print("Less than 2 tables")
			data = readDocx.getText(filename)
			me = data.replace("\\g","mmm")

			cell_data = []
			reactor_face = []
			span_list = []
			s_list = []
			if count_num == []:
				count = 6 
			else:
				count = max(count_num)
			for new_data in re.finditer("mmmEastmmm|mmmEast mmm|mmmEASTmmm|mmmEAST:mmm|mmmEAST mmm|mmmWestmmm|mmmWest mmm|mmmWESTmmm|mmmWEST:mmm|mmmWEST mmm", me):
				if new_data:
					span = new_data.span()
					span_list.append(span)
		
						#Reactor Face cell
					r_face = me[span[0]+3:span[1]-3] 
					reactor_face.append(r_face)
					s = list(range(len(span_list)))
					s_list.append(s)
				exit

			s_list2 = list(range(len(s_list)))
			sub_para = []
		
			for x in s_list2:
			
 
				if x == max(s_list2):
					count +=1
					sub_paragraphs = me[span_list[x][1]-3:len(me)]
					sub_para.append(str(sub_paragraphs)) 
									

				else:
					sub_paragraphs = me[span_list[x][1]-3:span_list[x+1][0]]
					sub_para.append(str(sub_paragraphs))
				exit	
		
		
			span2_list = []
			s2_list = []
			sub_para_str = str(sub_para)

			for sentence in re.finditer("m{3,}", sub_para_str):
				if sentence: 
					span2 = sentence.span()		
					span2_list.append(span2)
					s2 = list(range(len(span2_list)))
					s2_list.append(s2)
				exit
		
			s2_list2 = list(range(len(s2_list)))
			count_2 = 0		
			loc = []
			col_size = []
			sub_sentences = []
			for z in s2_list2:	
			
				if z == max(s2_list2):

					sub_sentence = sub_para_str[span2_list[z][1]:len(me)]
					sub_sentences.append(sub_sentence)
			
				else:
					sub_sentence = sub_para_str[span2_list[z][1]:span2_list[z+1][0]]
					sub_sentence = sub_sentence.replace(u'\\xa0', u'')
					sub_sentences.append(sub_sentence)

					if re.search("', '|', \"", sub_sentence):

							count += 1
							count_2 += 1						

					else:
						count += 1
						coun += 1
						loc_i = str('F'+str(count))
						loc_x = str('P'+str(count))
						loc_d = str('A'+str(count))
					
						worksheet.write(loc_i,reactor_face[count_2],cell_format_pop)
						worksheet.write(loc_x,sub_sentence,cell_format_info)
						worksheet.write(loc_d,coun,cell_format_info)

						length = len(sub_sentence)
						col_size.append(length)
						worksheet.set_column('P:P', max(col_size)-100)

						location =[]
						times = [] 
						for time in re.finditer("\d{1,2}:\d\dam|\d{1,2}:\d\dpm|\d{1,2}:\d\d|\d{1,2} oclock|\d{4}h|\d{1,2}noon", sub_sentence):
						
							if time:
								location.append(z)
								span3 = time.span()
								time = sub_sentence[span3[0]:span3[1]]	
								times.append(time)
								if len(location) == 2:
									loc_t = str('J'+str(count))
									worksheet.write(loc_t,time,cell_format_info)
							else:
								loc_t = str('I'+str(count))
								worksheet.write(loc_t,time,cell_format_info)
						exit

					location = []
					latice_sites = []
					for latice_site in re.finditer("[A-Z]{1}[0-9]{1,2} |[A-Z]{1}[0-9]{1,2}| [A-Z]{1}[0-9]{2} |[A-Z]{1}[0-9]{2},|[A-Z]{1}[0-9]{1},|[A-Z]{1}[0-9]{1,2}-|[A-Z]{1}[0-9]{1,2}-|[A-Z]{1}[0-9]{2}W|[A-Z]{1}[0-9]{1,2}E|[A-Z]{1}[0-9]{2}W:|[A-Z]{1}[0-9]{2}E:", sub_sentence):
	
						if latice_site:
							location.append(z)
							span4 = latice_site.span()
							latice_site = sub_sentence[span4[0]:span4[1]]
							latice_sites.append(latice_site)
							
							
							if len(location) == 1:
								main_sites.append(latice_site)
								loc_s = str('G'+str(count))
								worksheet.write(loc_s,latice_site,cell_format_info)
								#print(latice_site,"'First Site'")

							else:
								main = main_sites
								main.append(latice_site)
								cell_text = ','.join(main)
								loc_c = str('G'+str(count))
								worksheet.write(loc_c,cell_text,cell_format_info)

					
						for opn in re.finditer("opn\d+.\d+", sub_sentence):
							if opn:
								span = opn.span()
								opn = sub_sentence[span[0]:span[1]]	
								loc_o = str('H'+str(count))
								worksheet.write(loc_o,opn,cell_format_info)
						

						for crit_path in re.finditer("Critical path delay ~", sub_sentence):
							if crit_path:
								ans = 'Y'
								loc_c = str('K'+str(count))
								worksheet.write(loc_c,ans,cell_format_info)


				count_lists = []
				count_numb = [] 
				main_sites = []			
				for date in re.finditer(r"\w+\s\d+\s\d{4}|\w+\s\d+th|\w+\s\d+\sDay|\w+\s\d+\sday|\w+\s\d+\s(Day)|\w+\s\d+\s(day)|\w+\s\d+\sDAY|\w+\s\d+\sNight|\w+\s\d+\snight|\w+\s\d+\s(Night)|\w+\s\d+\s(night)|\w+\s\d+\sNIGHT|\w+\s\d[1,2],",sub_sentence):

					if date:
						span = date.span()
						

						for day in re.finditer(r"\w+\s\d+\s|\w+\s\d+,|\w+\s\d+t", sub_sentence):
							if day:
								span2 = day.span()
								date = sub_sentence[span2[0]:span2[1]-1]
								count_numb.append(count)
								loc_d = str('B'+str(count))
								worksheet.write(loc_d,date,cell_format_info)

						for year in re.finditer('\d{4}', sub_sentence):
							if year:
								span3 = year.span()
								year = sub_sentence[span3[0]:span3[1]]
								loc_y = str('D'+str(count)) 
								worksheet.write(loc_y,year,cell_format_info)

						for shift_n in re.finditer('Night|night|NIGHT', sub_sentence):
							if shift_n:
								shift = "N"
								loc_s = str('C'+str(count)) 
								worksheet.write(loc_s,shift,cell_format_info)					


						for shift_d in re.finditer('Day|day|DAY', sub_sentence):
							if shift_d:
								shift = "D"
								loc_s = str('C'+str(count)) 
								worksheet.write(loc_s,shift,cell_format_info)


				exit
	workbook.close()
	
	

			
